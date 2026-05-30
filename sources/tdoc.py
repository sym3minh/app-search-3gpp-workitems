"""
tdoc.py — Download và extract tài liệu TDoc.

Import: config, cr_search.download_cr_file (lazy, chỉ dùng trong tdoc_fetch_from_db)
Module phức tạp nhất — xử lý toàn bộ pipeline từ download đến extract text Word.
"""

import re, os, struct, sqlite3, datetime, time, zipfile
from pathlib import Path

from config import (
    TDOC_FETCH_OK, TDOC_DOCX_OK,
    PORTAL_BASE, HDRS,
    DB_FILE, CACHE_DIR,
    DOWNLOAD_ZIP_DIR, DOWNLOAD_EXTRACTED_DIR, DATA_DIR,
)
from cr_extractor import extract_cr_metadata
from heading_extractor import extract_headings_from_docx
from ts_info_db import TsInfoDb


# ── Exceptions ────────────────────────────────────────────────────────────────

class NoCRFound(Exception):
    """Portal trả về 'No Change Request found' hoặc DB không có row nào."""


class NoAgreedTDocs(Exception):
    """Có CR nhưng không có TDoc nào ở trạng thái 'agreed' / không có download URL hợp lệ."""


# ── Word namespace ─────────────────────────────────────────────────────────────
_WNS       = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
_WORD_EXTS = {'.docx', '.doc', '.dot', '.dotx', '.docm', '.dotm'}


# ══════════════════════════════════════════════════════════════════════════════
# ZIP / file cache helpers
# ══════════════════════════════════════════════════════════════════════════════

def _zip_cache_path(fname: str) -> Path:
    """
    Return canonical Zip cache path.
    Layout:  data/downloads/Zip/<FIRST_CHAR_UPPER>/<fname>
    """
    first_char = fname[0].upper() if fname else "_"
    if not first_char.isalnum():
        first_char = "_"
    folder = DOWNLOAD_ZIP_DIR / first_char
    folder.mkdir(parents=True, exist_ok=True)
    return folder / fname


def _find_zip_in_cache(stem: str) -> "Path | None":
    """Search all letter-subfolders under DOWNLOAD_ZIP_DIR for <stem>.zip."""
    fname      = f"{stem}.zip"
    first_char = fname[0].upper() if fname else "_"
    if not first_char.isalnum():
        first_char = "_"
    candidate = DOWNLOAD_ZIP_DIR / first_char / fname
    if candidate.exists() and candidate.stat().st_size > 500:
        return candidate
    if DOWNLOAD_ZIP_DIR.exists():
        for sub in DOWNLOAD_ZIP_DIR.iterdir():
            if sub.is_dir():
                p = sub / fname
                if p.exists() and p.stat().st_size > 500:
                    return p
    return None


def _extract_zip_to(zip_path: Path, extract_dir: Path, log_fn=None) -> bool:
    """Extract a ZIP file to extract_dir. Returns True on success."""
    def log(m):
        if log_fn:
            log_fn(m)
    try:
        extract_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'r') as z:
            z.extractall(extract_dir)
        log(f"Extracted {zip_path.name} → {extract_dir}")
        return True
    except Exception as e:
        log(f"[WARN] Zip extract lỗi {zip_path.name}: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# File type detection
# ══════════════════════════════════════════════════════════════════════════════

def _detect_ext(content: bytes, content_type: str = '', content_disposition: str = ''):
    """
    Detect file extension from:
    1. Content-Disposition filename  (highest priority)
    2. Magic bytes
    3. Content-Type  (fallback)

    Returns (ext: str, original_filename: str | None)
    """
    # 1. Content-Disposition filename
    if 'filename' in content_disposition:
        fm = re.search(r'filename[^;=\n]*=(([\'"]).*?\2|[^;\n]*)', content_disposition)
        if fm:
            name = fm.group(1).strip('\'"').strip()
            ext  = os.path.splitext(name)[1].lower()
            if ext in ('.zip', '.doc', '.docx', '.pdf', '.ppt', '.pptx', '.xls', '.xlsx'):
                return ext, name

    # 2. Magic bytes
    magic = content[:8] if len(content) >= 8 else content
    if magic[:4] == b'PK\x03\x04':
        return '.zip', None
    if magic[:4] in (b'\xd0\xcf\x11\xe0',):
        return '.doc', None
    if magic[:4] == b'%PDF':
        return '.pdf', None
    if magic[:2] == b'\x1f\x8b':
        return '.zip', None
    if magic[:4] == b'Rar!':
        return '.zip', None

    # 3. Content-Type
    ct = content_type.lower()
    if 'zip'         in ct: return '.zip',  None
    if 'officedoc'   in ct: return '.docx', None
    if 'msword'      in ct or 'word' in ct: return '.doc', None
    if 'pdf'         in ct: return '.pdf',  None
    return '.doc', None  # last resort


# ══════════════════════════════════════════════════════════════════════════════
# Single TDoc download
# ══════════════════════════════════════════════════════════════════════════════

def _tdoc_download_one(session, tdoc_info: dict, extract_dir: Path, log_fn=None) -> "Path | None":
    """
    Download a single TDoc file.

    ZIP  → saved to data/downloads/Zip/<LETTER>/<fname>.zip (shared cache).
           Extraction → extract_dir/<stem>/
    Non-ZIP → saved directly into extract_dir.

    Returns final Path of saved file, or None on failure.
    """
    from urllib.parse import urljoin, urlparse, parse_qs
    from bs4 import BeautifulSoup

    url = tdoc_info['download_url']
    num = tdoc_info['tdoc_number']
    if url.startswith('javascript:'):
        return None
    if not url.startswith('http'):
        url = urljoin('https://portal.3gpp.org/', url)
    try:
        resp    = session.get(url, timeout=60, verify=False)
        content = resp.content

        # HTML redirect page (e.g. window.location.href = "real_url")
        if (len(content) < 5000 and
                (b'<html' in content[:200].lower() or b'<!doctype' in content[:200].lower())):
            s = BeautifulSoup(content, 'html.parser')
            real_url = None
            for sc in s.find_all('script'):
                txt = sc.string or ''
                if 'window.location.href' in txt:
                    m = re.search(r"window\.location\.href\s*=\s*['\"]([^'\"]+)['\"]", txt)
                    if m:
                        real_url = m.group(1)
                        break
            if not real_url:
                return None
            resp    = session.get(real_url, timeout=60, verify=False)
            content = resp.content

        if len(content) < 200:
            return None
        if b'<html' in content[:200].lower():
            return None

        ct  = resp.headers.get('content-type', '')
        cd  = resp.headers.get('content-disposition', '')
        ext, cd_filename = _detect_ext(content, ct, cd)

        # Determine filename stem
        if cd_filename:
            stem, _ = os.path.splitext(cd_filename)
            fname   = cd_filename
        else:
            qs        = parse_qs(urlparse(url).query)
            uid_param = (qs.get('contributionUid') or qs.get('tdocuid') or [None])[0]
            stem      = uid_param if uid_param else num
            fname     = f"{stem}{ext}"

        if ext == '.zip':
            zip_path = _zip_cache_path(fname)
            if not (zip_path.exists() and zip_path.stat().st_size > 500):
                zip_path.write_bytes(content)
            xdir = extract_dir / stem
            if not xdir.exists():
                _extract_zip_to(zip_path, xdir, log_fn=log_fn)
            return zip_path
        else:
            extract_dir.mkdir(parents=True, exist_ok=True)
            fp = extract_dir / fname
            if fp.exists() and fp.stat().st_size > 500:
                return fp
            fp.write_bytes(content)
            return fp

    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# Portal HTML parser
# ══════════════════════════════════════════════════════════════════════════════

def _tdoc_parse_agreed(html: str) -> list:
    """Parse HTML portal, lọc TDocs có status == 'agreed'."""
    from bs4 import BeautifulSoup
    soup   = BeautifulSoup(html, "html.parser")
    result = []
    kw     = {'agreed', 'postponed', 'revised', 'withdrawn', 'not treated', 'noted', 'rejected', 'merged'}
    for link in soup.find_all('a', id='wgTdocDetailsLink'):
        num  = link.text.strip()
        href = link.get('href', '')
        if not num or not href:
            continue
        row    = link.find_parent('tr')
        status = ''
        if row:
            for cell in row.find_all('td'):
                t = cell.text.strip().lower()
                if t in kw:
                    status = t
                    break
        if status == 'agreed':
            result.append({'tdoc_number': num, 'download_url': href})
    return result


# ══════════════════════════════════════════════════════════════════════════════
# Public download APIs
# ══════════════════════════════════════════════════════════════════════════════

def tdoc_fetch_agreed(uid: str, log_fn=None, stop_event=None):
    """
    Tải tất cả 'agreed' TDocs của 1 work item.
    Scrape CR list nhiều trang, deduplicate, tải song song từng file.

    Raises NoCRFound, NoAgreedTDocs.
    Returns (downloaded, skipped, errors).
    """
    if not TDOC_FETCH_OK:
        raise RuntimeError("Thiếu thư viện:\npip install requests beautifulsoup4")

    import requests

    def log(m):
        if log_fn:
            log_fn(m)

    base = f"https://portal.3gpp.org/ChangeRequests.aspx?q=1&workitem={uid}"
    sess = requests.Session()
    sess.headers.update({
        'User-Agent': HDRS['User-Agent'],
        'Accept':     'text/html,application/xhtml+xml,*/*;q=0.8',
    })
    try:
        sess.get('https://portal.3gpp.org/', timeout=10, verify=False)
    except Exception:
        pass

    url = base + "&rgCrListChangePage=3_200"
    log(f"Fetching CR list for WI {uid}...")
    try:
        resp = sess.get(url, timeout=30, verify=False)
        resp.raise_for_status()
    except Exception as e:
        raise RuntimeError(f"Không tải được trang CR: {e}")

    if "No Change Request found" in resp.text:
        raise NoCRFound()

    total_pages = 1
    from bs4 import BeautifulSoup
    pager = BeautifulSoup(resp.text, 'html.parser').find('div', class_='rgInfoPart')
    if pager:
        m = re.search(r'(\d+)\s+pages', pager.text)
        if m:
            total_pages = int(m.group(1))
    log(f"Tổng {total_pages} trang CR")

    all_agreed = _tdoc_parse_agreed(resp.text)
    for pn in range(2, total_pages + 1):
        if stop_event and stop_event.is_set():
            break
        try:
            pr = sess.get(f"{base}&rgCrListChangePage={pn-1}_200", timeout=30, verify=False)
            all_agreed.extend(_tdoc_parse_agreed(pr.text))
        except Exception as e:
            log(f"[WARN] Trang {pn} lỗi: {e}")

    # Deduplicate
    seen = set(); deduped = []
    for t in all_agreed:
        if t['tdoc_number'] not in seen:
            seen.add(t['tdoc_number'])
            deduped.append(t)
    all_agreed = deduped

    if not all_agreed:
        raise NoAgreedTDocs()

    log(f"Tìm thấy {len(all_agreed)} agreed TDoc(s), đang tải...")
    extract_dir = DOWNLOAD_EXTRACTED_DIR / f"workitem_{uid}" / "agreed"
    extract_dir.mkdir(parents=True, exist_ok=True)

    downloaded = []; skipped = 0; errors = 0
    for i, tdoc in enumerate(all_agreed, 1):
        if stop_event and stop_event.is_set():
            log(f"⏹ Dừng theo yêu cầu sau {i-1}/{len(all_agreed)} file(s).")
            break
        num = tdoc['tdoc_number']

        cached_zip = _find_zip_in_cache(num)
        if cached_zip:
            log(f"[{i}/{len(all_agreed)}] CACHE {num} (zip cache hit)")
            xdir = extract_dir / num
            if not xdir.exists():
                _extract_zip_to(cached_zip, xdir, log_fn=log)
            skipped += 1
            downloaded.append(cached_zip)
            continue

        existing = next(
            (extract_dir / f"{num}{ext}" for ext in ['.doc', '.docx', '.pdf']
             if (extract_dir / f"{num}{ext}").exists()), None
        )
        if existing and existing.stat().st_size > 1000:
            log(f"[{i}/{len(all_agreed)}] SKIP {num} (exists)")
            skipped += 1
            downloaded.append(existing)
            continue

        fp = _tdoc_download_one(sess, tdoc, extract_dir)
        if fp:
            log(f"[{i}/{len(all_agreed)}] OK   {num}  ({fp.stat().st_size//1024} KB)")
            downloaded.append(fp)
        else:
            log(f"[{i}/{len(all_agreed)}] FAIL {num}")
            errors += 1
        time.sleep(0.3)

    return downloaded, skipped, errors


def tdoc_fetch_from_db(uid: str, log_fn=None):
    """
    Fallback: tải CR files từ cr_titles.db khi portal không trả về kết quả.
    Hỗ trợ cả schema mới (wg_tdoc/tsg_tdoc) và schema cũ (download_url).

    Raises NoCRFound, NoAgreedTDocs.
    Returns (downloaded, skipped, errors).
    """
    def log(m):
        if log_fn:
            log_fn(m)

    if not DB_FILE.exists():
        raise RuntimeError(f"Không có database tại {DB_FILE}")

    conn = sqlite3.connect(str(DB_FILE))
    cols = {row[1] for row in conn.execute("PRAGMA table_info(cr_titles)").fetchall()}
    has_wg  = "wg_tdoc"      in cols
    has_tsg = "tsg_tdoc"     in cols
    has_dl  = "download_url" in cols

    if has_wg or has_tsg:
        select_cols = "title"
        if has_wg:  select_cols += ", wg_tdoc"
        if has_tsg: select_cols += ", tsg_tdoc"
        where = " OR ".join(
            [f"{c} IS NOT NULL AND {c} != ''"
             for c in (["wg_tdoc"] if has_wg else []) + (["tsg_tdoc"] if has_tsg else [])]
        )
        raw = conn.execute(
            f"SELECT {select_cols} FROM cr_titles WHERE workitem_id=? AND ({where})",
            (uid,)
        ).fetchall()
        conn.close()
        if not raw:
            raise NoCRFound()
        valid = []
        for row in raw:
            title  = row[0]
            wg     = row[1].strip() if has_wg  and len(row) > 1 and row[1] else ""
            tsg    = row[2].strip() if has_tsg and len(row) > 2 and row[2] else (
                     row[1].strip() if has_tsg and not has_wg and len(row) > 1 and row[1] else "")
            tdoc_id = wg if wg else tsg
            if tdoc_id:
                valid.append((title, tdoc_id))
    elif has_dl:
        raw = conn.execute(
            "SELECT title, download_url FROM cr_titles "
            "WHERE workitem_id=? AND download_url IS NOT NULL AND download_url != ''",
            (uid,)
        ).fetchall()
        conn.close()
        if not raw:
            raise NoCRFound()
        valid = [(t, u) for t, u in raw
                 if u and not u.strip().startswith('javascript:')]
    else:
        conn.close()
        raise RuntimeError("cr_titles không có cột wg_tdoc/tsg_tdoc hoặc download_url")

    if not valid:
        raise NoAgreedTDocs()

    log(f"Tìm thấy {len(valid)} CR(s) trong database cho WI {uid}, đang tải...")
    extract_dir = DOWNLOAD_EXTRACTED_DIR / f"workitem_{uid}" / "db_crs"
    extract_dir.mkdir(parents=True, exist_ok=True)

    if not TDOC_FETCH_OK:
        raise RuntimeError("Thiếu thư viện:\npip install requests beautifulsoup4")

    # Import lazily to avoid circular import at module level
    from cr_search import download_cr_file

    import requests
    downloaded = []; skipped = 0; errors = 0
    for i, (title, tdoc_or_url) in enumerate(valid, 1):
        if tdoc_or_url.startswith('http'):
            from urllib.parse import urlparse, parse_qs
            qs        = parse_qs(urlparse(tdoc_or_url).query)
            uid_param = (qs.get('contributionUid') or qs.get('tdocuid') or [None])[0]
            stem      = uid_param or re.sub(r'[^\w\-]', '_', title[:60]).strip('_') or f"cr_{i}"
            dl_url    = tdoc_or_url
        else:
            stem   = tdoc_or_url
            dl_url = (f"https://portal.3gpp.org/ngppapp/DownloadTDoc.aspx"
                      f"?contributionUid={stem}")

        cached_zip = _find_zip_in_cache(stem)
        if cached_zip:
            log(f"[{i}/{len(valid)}] CACHE {stem} (zip cache hit)")
            xdir = extract_dir / stem
            if not xdir.exists():
                _extract_zip_to(cached_zip, xdir, log_fn=log_fn)
            skipped += 1
            downloaded.append(cached_zip)
            continue

        existing = next(
            (extract_dir / f"{stem}{ext}" for ext in ['.doc', '.docx', '.pdf']
             if (extract_dir / f"{stem}{ext}").exists()), None
        )
        if existing and existing.stat().st_size > 500:
            log(f"[{i}/{len(valid)}] SKIP {stem} (exists)")
            skipped += 1
            downloaded.append(existing)
            continue

        fp = download_cr_file(dl_url, extract_dir, hint_name=title)
        if fp:
            log(f"[{i}/{len(valid)}] OK   {fp.name}  ({fp.stat().st_size//1024} KB)")
            downloaded.append(fp)
        else:
            log(f"[{i}/{len(valid)}] FAIL {stem}")
            errors += 1
        time.sleep(0.2)

    return downloaded, skipped, errors


# ══════════════════════════════════════════════════════════════════════════════
# Smart fetch — DB-first, then portal incremental sync
# ══════════════════════════════════════════════════════════════════════════════




def _db_get_wi_cr_count(uid: str) -> int:
    """Trả về số CR titles của WI trong cr_titles.db. 0 nếu DB không tồn tại."""
    if not DB_FILE.exists():
        return 0
    try:
        conn = sqlite3.connect(str(DB_FILE))
        n = conn.execute(
            "SELECT COUNT(*) FROM cr_titles WHERE workitem_id=?", (uid,)
        ).fetchone()[0]
        conn.close()
        return n
    except Exception:
        return 0


def _db_get_first_tdoc_id(uid: str) -> str:
    """
    Lấy first_tdoc_id từ bảng workitems trong cr_titles.db.
    Trả về chuỗi rỗng nếu không có.
    """
    if not DB_FILE.exists():
        return ""
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cols = {row[1] for row in conn.execute("PRAGMA table_info(workitems)").fetchall()}
        if "first_tdoc_id" not in cols:
            conn.close()
            return ""
        row = conn.execute(
            "SELECT first_tdoc_id FROM workitems WHERE workitem_id=?", (uid,)
        ).fetchone()
        conn.close()
        return (row[0] or "").strip() if row else ""
    except Exception:
        return ""


def _db_upsert_cr(uid: str, title: str, wg_tdoc: str, tsg_tdoc: str,
                  status: str, release: str = "", log_fn=None):
    """
    Upsert 1 CR record vào cr_titles và cập nhật workitems:
      - last_crawled
      - cr_count  (tăng thêm 1)
      - first_tdoc_id  (ghi đè bằng wg_tdoc mới nhất được thấy đầu tiên khi sync)
    Chỉ insert nếu wg_tdoc chưa tồn tại trong DB.
    Trả về True nếu đã insert mới, False nếu đã có sẵn.
    """
    def log(m):
        if log_fn: log_fn(m)
    if not DB_FILE.exists():
        return False
    try:
        conn = sqlite3.connect(str(DB_FILE))
        cr_cols = {row[1] for row in conn.execute("PRAGMA table_info(cr_titles)").fetchall()}
        has_wg      = "wg_tdoc"   in cr_cols
        has_tsg     = "tsg_tdoc"  in cr_cols
        has_release = "release"   in cr_cols

        # Kiểm tra đã có chưa
        if has_wg and wg_tdoc:
            existing = conn.execute(
                "SELECT id FROM cr_titles WHERE workitem_id=? AND wg_tdoc=?",
                (uid, wg_tdoc)
            ).fetchone()
            if existing:
                conn.close()
                return False

        # Build INSERT động theo schema thực tế
        insert_cols  = ["title", "workitem_id"]
        insert_vals  = [title, uid]
        if has_wg:
            insert_cols.append("wg_tdoc");  insert_vals.append(wg_tdoc)
        if has_tsg:
            insert_cols.append("tsg_tdoc"); insert_vals.append(tsg_tdoc)
        if has_release and release:
            insert_cols.append("release");  insert_vals.append(release)

        placeholders = ", ".join("?" * len(insert_cols))
        conn.execute(
            f"INSERT INTO cr_titles ({', '.join(insert_cols)}) VALUES ({placeholders})",
            insert_vals
        )

        # Cập nhật workitems
        wi_cols = {row[1] for row in conn.execute("PRAGMA table_info(workitems)").fetchall()}
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        set_parts  = []
        set_values = []

        if "last_crawled" in wi_cols:
            set_parts.append("last_crawled = ?")
            set_values.append(now)

        if "cr_count" in wi_cols:
            set_parts.append("cr_count = COALESCE(cr_count, 0) + 1")

        # first_tdoc_id: chỉ update khi cột tồn tại và WG tdoc hợp lệ
        # Giá trị được set bởi caller (_portal_incremental_sync) khi gặp row đầu tiên
        # Ở đây không tự động update để tránh ghi đè sai thứ tự

        if set_parts:
            conn.execute(
                f"UPDATE workitems SET {', '.join(set_parts)} WHERE workitem_id=?",
                set_values + [uid]
            )

        conn.commit()
        conn.close()
        log(f"  [DB] Inserted CR: {wg_tdoc} / {tsg_tdoc} release={release!r} (status={status})")
        return True
    except Exception as e:
        log(f"  [DB] Upsert lỗi: {e}")
        return False


def _portal_parse_cr_rows(html: str) -> list:
    """
    Parse toàn bộ CR rows từ trang portal (không lọc status).
    Trả về list dict: {tdoc_number, tsg_tdoc, release, status, title, href}
    """
    from bs4 import BeautifulSoup
    soup   = BeautifulSoup(html, "html.parser")
    result = []
    all_statuses = {
        'agreed', 'approved', 'postponed', 'revised', 'withdrawn',
        'not treated', 'noted', 'rejected', 'merged'
    }
    for link in soup.find_all('a', id='wgTdocDetailsLink'):
        wg_num = link.text.strip()
        href   = link.get('href', '')
        if not wg_num:
            continue
        row = link.find_parent('tr')
        if not row:
            continue
        cells = row.find_all('td')
        status   = ''
        tsg_tdoc = ''
        title    = ''
        release  = ''
        for cell in cells:
            t  = cell.text.strip()
            tl = t.lower()
            if tl in all_statuses and not status:
                status = tl
            # TSG TDoc thường có dạng RP-xxxxxx
            if re.match(r'^RP-\d+$', t) and not tsg_tdoc:
                tsg_tdoc = t
            # Release: link openRelease hoặc text dạng Rel-NN
            if not release:
                rel_link = cell.find('a', href=re.compile(r'openRelease', re.I))
                if rel_link:
                    release = rel_link.text.strip()
                elif re.match(r'^Rel-\d+$', t, re.I):
                    release = t
        # Lấy title từ cột Subject: text dài, không phải TDoc ID hay Release
        for cell in cells:
            txt = cell.text.strip()
            if (len(txt) > 20
                    and not re.match(r'^(RP-|R\d+-)\d+', txt)
                    and not re.match(r'^Rel-\d+$', txt, re.I)):
                title = txt
                break
        result.append({
            'tdoc_number': wg_num,
            'tsg_tdoc':    tsg_tdoc,
            'release':     release,
            'status':      status,
            'title':       title,
            'href':        href,
        })
    return result


def _portal_incremental_sync(uid: str, first_tdoc_id: str, sess,
                              extract_dir: Path, log_fn=None) -> tuple:
    """
    Fetch portal CR list và:
      - Duyệt từng row từ đầu trang (mới nhất trước)
      - Row đầu tiên gặp được → cập nhật first_tdoc_id trên workitems (nếu cột tồn tại)
      - Nếu wg_tdoc == first_tdoc_id → dừng (đã đến mốc đã biết)
      - status == 'approved' → upsert DB (cr_titles + workitems.cr_count++)
      - status == 'agreed'   → upsert DB + download ngay

    Sau khi duyệt xong: đồng bộ lại workitems.cr_count = COUNT(*) thực tế trong cr_titles.

    Trả về (downloaded, skipped, errors, new_count).
    """
    from bs4 import BeautifulSoup

    def log(m):
        if log_fn: log_fn(m)

    downloaded = []; skipped = 0; errors = 0; new_count = 0
    base = f"https://portal.3gpp.org/ChangeRequests.aspx?q=1&workitem={uid}"

    try:
        resp = sess.get(base + "&rgCrListChangePage=3_200", timeout=30, verify=False)
        resp.raise_for_status()
    except Exception as e:
        log(f"  [Sync] Không fetch được portal: {e}")
        return downloaded, skipped, errors, new_count

    if "No Change Request found" in resp.text:
        log(f"  [Sync] Portal: No CR found for WI {uid}")
        return downloaded, skipped, errors, new_count

    total_pages = 1
    pager = BeautifulSoup(resp.text, 'html.parser').find('div', class_='rgInfoPart')
    if pager:
        m = re.search(r'(\d+)\s+pages', pager.text)
        if m:
            total_pages = int(m.group(1))
    log(f"  [Sync] Portal: {total_pages} trang CR cho WI {uid}")

    pages_html = [resp.text]
    for pn in range(2, total_pages + 1):
        try:
            pr = sess.get(f"{base}&rgCrListChangePage={pn-1}_200", timeout=30, verify=False)
            pages_html.append(pr.text)
        except Exception as e:
            log(f"  [Sync] Trang {pn} lỗi: {e}")

    stop_flag        = False
    first_seen_tdoc  = None   # wg_tdoc đầu tiên nhìn thấy trên portal

    for page_html in pages_html:
        if stop_flag:
            break
        rows = _portal_parse_cr_rows(page_html)
        for row in rows:
            wg_num  = row['tdoc_number']
            tsg_num = row['tsg_tdoc']
            release = row['release']
            status  = row['status']
            title   = row['title']
            href    = row['href']

            # Ghi nhớ WG TDoc đầu tiên nhìn thấy (mới nhất) để update first_tdoc_id
            if first_seen_tdoc is None:
                first_seen_tdoc = wg_num

            # Gặp mốc first_tdoc_id → dừng scan
            if first_tdoc_id and wg_num == first_tdoc_id:
                log(f"  [Sync] Gặp first_tdoc_id={first_tdoc_id} → dừng scan")
                stop_flag = True
                break

            if status == 'approved':
                inserted = _db_upsert_cr(uid, title, wg_num, tsg_num, status,
                                         release=release, log_fn=log_fn)
                if inserted:
                    new_count += 1

            elif status == 'agreed':
                inserted = _db_upsert_cr(uid, title, wg_num, tsg_num, status,
                                         release=release, log_fn=log_fn)
                if inserted:
                    new_count += 1
                # Download ngay
                dl_url = (href if href and href.startswith('http')
                          else f"https://portal.3gpp.org/ngppapp/DownloadTDoc.aspx"
                               f"?contributionUid={wg_num}")
                cached = _find_zip_in_cache(wg_num)
                if cached:
                    log(f"  [Sync] CACHE {wg_num}")
                    xdir = extract_dir / wg_num
                    if not xdir.exists():
                        _extract_zip_to(cached, xdir, log_fn=log_fn)
                    skipped += 1
                    downloaded.append(cached)
                    continue
                existing = next(
                    (extract_dir / f"{wg_num}{ext}" for ext in ['.doc', '.docx', '.pdf']
                     if (extract_dir / f"{wg_num}{ext}").exists()), None
                )
                if existing and existing.stat().st_size > 500:
                    log(f"  [Sync] SKIP {wg_num} (exists)")
                    skipped += 1; downloaded.append(existing); continue

                tdoc_info = {'tdoc_number': wg_num, 'download_url': dl_url}
                fp = _tdoc_download_one(sess, tdoc_info, extract_dir, log_fn=log_fn)
                if fp:
                    log(f"  [Sync] OK {wg_num} ({fp.stat().st_size // 1024} KB)")
                    downloaded.append(fp)
                else:
                    log(f"  [Sync] FAIL {wg_num}")
                    errors += 1
                time.sleep(0.3)

    # ── Cập nhật workitems sau khi sync ──────────────────────────────────────
    if DB_FILE.exists():
        try:
            conn = sqlite3.connect(str(DB_FILE))
            wi_cols = {row[1] for row in conn.execute("PRAGMA table_info(workitems)").fetchall()}

            set_parts  = []
            set_values = []
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            if "last_crawled" in wi_cols:
                set_parts.append("last_crawled = ?")
                set_values.append(now)

            # Đồng bộ cr_count = số thực tế trong cr_titles
            if "cr_count" in wi_cols:
                actual_count = conn.execute(
                    "SELECT COUNT(*) FROM cr_titles WHERE workitem_id=?", (uid,)
                ).fetchone()[0]
                set_parts.append("cr_count = ?")
                set_values.append(actual_count)
                log(f"  [Sync] workitems.cr_count → {actual_count}")

            # Cập nhật first_tdoc_id bằng WG TDoc mới nhất vừa thấy
            if "first_tdoc_id" in wi_cols and first_seen_tdoc:
                set_parts.append("first_tdoc_id = ?")
                set_values.append(first_seen_tdoc)
                log(f"  [Sync] workitems.first_tdoc_id → {first_seen_tdoc}")

            if set_parts:
                conn.execute(
                    f"UPDATE workitems SET {', '.join(set_parts)} WHERE workitem_id=?",
                    set_values + [uid]
                )
                conn.commit()
            conn.close()
        except Exception as e:
            log(f"  [Sync] Lỗi cập nhật workitems: {e}")

    log(f"  [Sync] Xong: {new_count} CR mới, {len(downloaded)} đã tải, {errors} lỗi")
    return downloaded, skipped, errors, new_count


def _extract_wg_from_tsg_zip(tsg_zip: Path, wg_id: str,
                               extract_dir: Path, log_fn=None) -> "Path | None":
    """
    Từ file zip của TSG TDoc, tìm và extract file Word có tên bắt đầu bằng wg_id.
    Ví dụ: tsg_zip=RP-191267.zip, wg_id='R4-1905981'
           → tìm file 'R4-1905981*.doc' hoặc 'R4-1905981*.docx' bên trong zip
    Trả về Path của file đã extract, hoặc None nếu không tìm thấy.
    """
    def log(m):
        if log_fn: log_fn(m)

    if not tsg_zip or not tsg_zip.exists():
        return None
    try:
        with zipfile.ZipFile(tsg_zip, 'r') as z:
            names = z.namelist()
            # Tìm entry bắt đầu bằng wg_id (case-insensitive)
            wg_lower = wg_id.lower()
            target = next(
                (n for n in names
                 if Path(n).name.lower().startswith(wg_lower)
                 and Path(n).suffix.lower() in _WORD_EXTS),
                None
            )
            if not target:
                log(f"  [TSG-extract] Không tìm thấy file bắt đầu bằng '{wg_id}' trong {tsg_zip.name}")
                log(f"  [TSG-extract] Các file trong zip: {[Path(n).name for n in names if Path(n).suffix.lower() in _WORD_EXTS]}")
                return None
            extract_dir.mkdir(parents=True, exist_ok=True)
            out_path = extract_dir / Path(target).name
            if out_path.exists() and out_path.stat().st_size > 500:
                log(f"  [TSG-extract] SKIP {out_path.name} (exists)")
                return out_path
            z.extract(target, extract_dir)
            # zipfile extract giữ nguyên subpath, ta cần flatten
            extracted = extract_dir / target
            if extracted != out_path and extracted.exists():
                import shutil
                shutil.move(str(extracted), str(out_path))
            log(f"  [TSG-extract] OK {out_path.name}")
            return out_path
    except Exception as e:
        log(f"  [TSG-extract] Lỗi: {e}")
        return None


def tdoc_fetch_smart(uid: str, log_fn=None, stop_event=None) -> tuple:
    """
    Chiến lược download TDoc cho Tab 1 và Tab 4.
    Luôn dùng WG TDoc trực tiếp (không dùng TSG path).

    Bước 1 — Đọc DB:
      - Lấy first_tdoc_id từ bảng workitems

    Bước 2 — Download từ DB (chỉ theo wg_tdoc):
      - Tầng 1: kiểm tra zip cache (data/downloads/Zip/...)
      - Tầng 2: kiểm tra extracted file đã có (data/downloads/Extracted/workitem_<uid>/...)
      - Nếu cả hai đều không có → download mới

    Bước 3 — Portal incremental sync:
      - Fetch CR list từ portal
      - Duyệt từng row, dừng khi gặp first_tdoc_id
      - approved → upsert DB; agreed → upsert DB + download

    Raises NoCRFound, NoAgreedTDocs.
    Returns (downloaded, skipped, errors).
    """
    if not TDOC_FETCH_OK:
        raise RuntimeError("Thiếu thư viện:\npip install requests beautifulsoup4")

    import requests as _req
    import urllib3 as _u3
    _u3.disable_warnings(_u3.exceptions.InsecureRequestWarning)

    def log(m):
        if log_fn: log_fn(m)

    # ── Bước 1: Đọc thông tin từ DB ──────────────────────────────────────────
    cr_count   = _db_get_wi_cr_count(uid)
    first_tdoc = _db_get_first_tdoc_id(uid)

    log(f"WI {uid}: {cr_count} CR titles trong DB, first_tdoc_id={first_tdoc!r}")

    extract_dir = DOWNLOAD_EXTRACTED_DIR / f"workitem_{uid}"
    extract_dir.mkdir(parents=True, exist_ok=True)

    sess = _req.Session()
    sess.headers.update({
        'User-Agent': HDRS['User-Agent'],
        'Accept':     'text/html,application/xhtml+xml,*/*;q=0.8',
    })
    try:
        sess.get('https://portal.3gpp.org/', timeout=10, verify=False)
    except Exception:
        pass

    all_downloaded = []; total_skipped = 0; total_errors = 0

    # ── Bước 2: Download từ DB (chỉ theo wg_tdoc) ────────────────────────────
    if not DB_FILE.exists():
        log(f"Không có DB tại {DB_FILE} — bỏ qua bước DB")
    else:
        conn = sqlite3.connect(str(DB_FILE))
        cols = {row[1] for row in conn.execute("PRAGMA table_info(cr_titles)").fetchall()}

        if "wg_tdoc" not in cols:
            log("DB không có cột wg_tdoc — bỏ qua bước DB")
            conn.close()
        else:
            rows = conn.execute(
                "SELECT title, wg_tdoc FROM cr_titles "
                "WHERE workitem_id=? AND wg_tdoc IS NOT NULL AND wg_tdoc != ''",
                (uid,)
            ).fetchall()
            conn.close()

            log(f"DB: {len(rows)} CR(s) để tải (direct WG)")

            for i, (title, wg_tdoc) in enumerate(rows, 1):
                if stop_event and stop_event.is_set():
                    log(f"⏹ Dừng theo yêu cầu sau {i-1} file(s).")
                    break

                dl_url = (f"https://portal.3gpp.org/ngppapp/DownloadTDoc.aspx"
                          f"?contributionUid={wg_tdoc}")

                # Tầng 1: zip cache check
                cached_zip = _find_zip_in_cache(wg_tdoc)
                if cached_zip:
                    log(f"[{i}/{len(rows)}] CACHE {wg_tdoc}")
                    xdir = extract_dir / wg_tdoc
                    if not xdir.exists():
                        _extract_zip_to(cached_zip, xdir, log_fn=log_fn)
                    all_downloaded.append(cached_zip)
                    total_skipped += 1
                    continue

                # Tầng 2: extracted file check
                existing = next(
                    (extract_dir / f"{wg_tdoc}{ext}"
                     for ext in ('.docx', '.doc', '.pdf')
                     if (extract_dir / f"{wg_tdoc}{ext}").exists()), None
                )
                if existing and existing.stat().st_size > 500:
                    log(f"[{i}/{len(rows)}] SKIP {wg_tdoc} (extracted exists)")
                    all_downloaded.append(existing)
                    total_skipped += 1
                    continue

                # Download mới
                tdoc_info = {'tdoc_number': wg_tdoc, 'download_url': dl_url}
                fp = _tdoc_download_one(sess, tdoc_info, extract_dir, log_fn=log_fn)
                if fp:
                    log(f"[{i}/{len(rows)}] OK {wg_tdoc} ({fp.stat().st_size // 1024} KB)")
                    all_downloaded.append(fp)
                else:
                    log(f"[{i}/{len(rows)}] FAIL {wg_tdoc}")
                    total_errors += 1
                time.sleep(0.25)

    # ── Bước 3: Portal incremental sync ──────────────────────────────────────
    if not (stop_event and stop_event.is_set()):
        log(f"Portal incremental sync cho WI {uid} (first_tdoc={first_tdoc!r})...")
        sync_dl, sync_skip, sync_err, new_count = _portal_incremental_sync(
            uid, first_tdoc, sess, extract_dir, log_fn=log_fn
        )
        all_downloaded.extend(sync_dl)
        total_skipped += sync_skip
        total_errors  += sync_err
        if new_count:
            log(f"Sync: {new_count} CR mới từ portal")

    if not all_downloaded and total_errors == 0:
        raise NoAgreedTDocs()

    return all_downloaded, total_skipped, total_errors, extract_dir


# ══════════════════════════════════════════════════════════════════════════════
# Word extraction helpers
# ══════════════════════════════════════════════════════════════════════════════

def _para_text_with_revisions(para):
    """Paragraph text: include track-change insertions, skip deletions."""
    try:
        elem  = para._element
        parts = []
        for ins in elem.findall(f'.//{_WNS}ins'):
            for t in ins.findall(f'.//{_WNS}t'):
                if t.text:
                    parts.append(t.text)
        for run in para.runs:
            ptag = run._element.getparent().tag if run._element.getparent() is not None else ''
            if f'{_WNS}ins' not in ptag and f'{_WNS}del' not in ptag:
                if run.text:
                    parts.append(run.text)
        return ''.join(parts).strip()
    except Exception:
        return (para.text or '').strip()


def _cell_text_with_revisions(cell):
    """Cell text: include track-change insertions, skip deletions."""
    try:
        tc    = cell._element
        parts = []
        for ins in tc.findall(f'.//{_WNS}ins'):
            for t in ins.findall(f'.//{_WNS}t'):
                if t.text:
                    parts.append(t.text)
        for para in cell.paragraphs:
            for run in para.runs:
                ptag = run._element.getparent().tag if run._element.getparent() is not None else ''
                if f'{_WNS}ins' not in ptag and f'{_WNS}del' not in ptag:
                    if run.text:
                        parts.append(run.text)
        return ' '.join(parts).strip()
    except Exception:
        return (cell.text or '').strip()


def _extract_table_with_merges(table):
    try:
        nr   = len(table.rows)
        nc   = len(table.columns)
        grid = [[None] * nc for _ in range(nr)]
        vmerge: dict = {}
        for i, row in enumerate(table.rows):
            ci = 0
            for cell in row.cells:
                while ci < nc and grid[i][ci] is not None:
                    ci += 1
                if ci >= nc:
                    break
                tcPr = cell._element.find(f'.//{_WNS}tcPr')
                span = 1
                if tcPr is not None:
                    gs = tcPr.find(f'.//{_WNS}gridSpan')
                    if gs is not None:
                        span = int(gs.get(f'{_WNS}val', '1'))
                vm_val = None
                if tcPr is not None:
                    vm = tcPr.find(f'.//{_WNS}vMerge')
                    if vm is not None:
                        vm_val = vm.get(f'{_WNS}val', 'continue')
                ctext = _cell_text_with_revisions(cell)
                if vm_val == 'restart':
                    vmerge[ci] = ctext; grid[i][ci] = ctext
                elif vm_val is not None or (vm_val is None and ci in vmerge and ctext == ''):
                    grid[i][ci] = vmerge.get(ci, '')
                else:
                    grid[i][ci] = ctext; vmerge.pop(ci, None)
                for s in range(1, span):
                    if ci + s < nc:
                        grid[i][ci + s] = ''
                ci += span
        return grid
    except Exception:
        return [[_cell_text_with_revisions(c) for c in row.cells] for row in table.rows]


def _table_to_markdown(table, table_num: int, caption=None) -> str:
    try:
        if not table.rows:
            return f'\n[Table {table_num}: empty]\n'
        grid = _extract_table_with_merges(table)
        if not any(any(str(c).strip() for c in r if c is not None) for r in grid):
            return f'\n[Table {table_num}: empty/template]\n'
        hdr_label = (
            f'\n**[Table {table_num}]: {caption}**\n' if caption
            else f'\n**[Table {table_num}]**\n'
        )
        lines = [hdr_label]
        hdr   = [str(c or '') for c in grid[0]]
        lines.append('| ' + ' | '.join(hdr) + ' |')
        lines.append('|' + '|'.join(['---'] * len(hdr)) + '|')
        for row in grid[1:]:
            if any(str(c).strip() for c in row if c is not None):
                lines.append('| ' + ' | '.join(str(c or '') for c in row) + ' |')
        lines.append('')
        return '\n'.join(lines)
    except Exception as e:
        return f'\n[Table {table_num}: error — {e}]\n'


def _detect_table_caption(last_para_text, table_num):
    if not last_para_text:
        return None
    for pat in [
        r'^Table\s+([\d.]+[-\w]*)\s*[:\-]?\s*(.+?)(?:\s+is\s+defined|\s+in\s+3GPP|$)',
        r'^Table\s+([\d.]+[-\w]*)\s+(.+?)(?:\s+is\s+defined|\s+in\s+3GPP|$)',
    ]:
        m = re.match(pat, last_para_text, re.IGNORECASE)
        if m:
            ref, title = m.group(1), m.group(2).strip()
            for phrase in ['is defined as', 'in 3GPP', 'clause', 'except for']:
                idx = title.lower().find(phrase)
                if idx > 0:
                    title = title[:idx].strip()
            return f'Table {ref}: {title}'
    kw = ['maximum', 'minimum', 'parameters', 'values', 'requirements', 'conditions',
          'configuration', 'band', 'frequency', 'power', 'reduction', 'class',
          'density', 'emission']
    if (len(last_para_text) < 150
            and not last_para_text.endswith('.')
            and any(k in last_para_text.lower() for k in kw)):
        return last_para_text.strip()
    return None


def _extract_3gpp_metadata(doc) -> dict:
    meta = {}
    try:
        if doc.paragraphs:
            first = doc.paragraphs[0].text
            m = re.search(r'([RSC]\d+-\d{6,7})', first)
            if m:
                meta['tdoc_number'] = m.group(1)
            m2 = re.search(r'(TSG-[^\t]+Meeting\s+#\d+)', first)
            if m2:
                meta['meeting'] = m2.group(1)
            if len(doc.paragraphs) > 1:
                meta['meeting_location'] = doc.paragraphs[1].text.strip()
        if len(doc.tables) > 0:
            rows = doc.tables[0].rows
            if len(rows) > 3:
                cells = [_cell_text_with_revisions(c) for c in rows[3].cells]
                if len(cells) > 1 and cells[1]:
                    meta['spec_number'] = cells[1].strip()
                if len(cells) > 3 and cells[3].strip().isdigit():
                    meta['cr_number'] = cells[3].strip()
                if len(cells) > 7 and cells[7]:
                    meta['spec_version'] = cells[7].strip()
        if len(doc.tables) > 2:
            for row in doc.tables[2].rows:
                cells = [_cell_text_with_revisions(c) for c in row.cells]
                if not cells or not cells[0]:
                    continue
                label = cells[0].strip().lower()
                val   = cells[1].strip() if len(cells) > 1 else ''
                val2  = cells[2].strip() if len(cells) > 2 else ''
                if   'title'             in label and val:  meta['title']             = val
                elif 'source to wg'      in label and val:  meta['source_wg']         = val
                elif 'source to tsg'     in label and val:  meta['source_tsg']        = val
                elif 'work item'         in label and val:  meta['work_item']         = val
                elif 'category'          in label and val:  meta['category']          = val
                elif 'release'           in label and val:  meta['release']           = val
                elif 'reason for change' in label and val2: meta['reason_for_change'] = val2
                elif 'summary of change' in label and val2: meta['summary_of_change'] = val2
    except Exception:
        pass
    return meta


# ── OLE2 binary .doc extraction ───────────────────────────────────────────────

def _clean_word_text(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = text.replace('\x07', '\t')
    text = text.replace('\x0B', '\n')
    text = text.replace('\x0C', '\n')
    text = re.sub(r'[^\x09\x0A\x20-\x7E\u00A0-\u024F\u1E00-\u1EFF]', ' ', text)
    text = re.sub(r'[ \t]{4,}', ' ', text)
    text = re.sub(r'\n{4,}', '\n\n', text)
    lines = [l.strip() for l in text.split('\n')]
    return '\n'.join(l for l in lines if l)


def _scan_word_stream_for_text(data: bytes) -> str:
    """Heuristic: find long runs of printable CP1252 text in the WordDocument stream."""
    try:
        decoded = data.decode('cp1252', errors='replace')
    except Exception:
        decoded = data.decode('latin-1', errors='replace')
    runs = re.findall(r'[\x20-\x7E\u00A0-\u024F]{4,}', decoded)
    good = []
    for run in runs:
        alpha = sum(1 for c in run if c.isalpha() or c.isspace())
        if alpha / max(len(run), 1) > 0.5:
            good.append(run.strip())
    return '\n'.join(good)


def _extract_ole2_printable_fallback(fp: Path) -> "tuple[str, str]":
    """Last-resort: scan raw bytes for CP1252 printable text runs (no olefile)."""
    try:
        raw     = fp.read_bytes()
        decoded = raw.decode('cp1252', errors='replace')
        runs    = re.findall(r'[\x20-\x7E\u00A0-\u024F]{6,}', decoded)
        good    = [r.strip() for r in runs
                   if sum(1 for c in r if c.isalpha() or c.isspace()) / max(len(r), 1) > 0.55]
        text    = '\n'.join(good)
        return _clean_word_text(text), 'raw CP1252 scan (install olefile for better results)'
    except Exception as e:
        return '', f'fallback error: {e}'


def _extract_ole2_doc_text(fp: Path) -> "tuple[str, str]":
    """
    Extract plain text from a true OLE2 binary .doc file via olefile.
    Returns (text, note).
    """
    try:
        import olefile
    except ImportError:
        return _extract_ole2_printable_fallback(fp)

    try:
        ole = olefile.OleFileIO(str(fp))
        if not ole.exists('WordDocument'):
            ole.close()
            return '', 'no WordDocument stream'
        wdoc = ole.openstream('WordDocument').read()
        ole.close()

        if len(wdoc) < 0x60:
            return '', 'WordDocument stream too short'

        flags    = struct.unpack_from('<H', wdoc, 10)[0]
        fc_min   = struct.unpack_from('<I', wdoc, 0x18)[0]
        ccp_text = struct.unpack_from('<I', wdoc, 0x4C)[0]

        if flags & 0x0100:
            return '', 'file is encrypted'
        if ccp_text == 0 or ccp_text > 5_000_000:
            return '', 'unexpected ccpText value'

        is_ansi    = bool(fc_min & 0x40000000)
        fc_start   = fc_min & 0x3FFFFFFF
        char_size  = 1 if is_ansi else 2
        byte_count = ccp_text * char_size

        if fc_start + byte_count <= len(wdoc):
            raw = wdoc[fc_start: fc_start + byte_count]
            text = raw.decode('cp1252', errors='replace') if is_ansi else raw.decode('utf-16-le', errors='replace')
        else:
            text = _scan_word_stream_for_text(wdoc)
            if not text:
                return '', 'text offset out of stream bounds'

        note = f"{'ANSI/CP1252' if is_ansi else 'UTF-16-LE'} via OLE2 stream"
        return _clean_word_text(text), note

    except Exception as e:
        return '', f'olefile error: {e}'


def _extract_ooxml_doc_text(fp: Path) -> str:
    """Some .doc files are actually OOXML ZIPs. Extract via word/document.xml."""
    try:
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(str(fp), 'r') as z:
            if 'word/document.xml' not in z.namelist():
                return ''
            xml_bytes = z.read('word/document.xml')
        root = ET.fromstring(xml_bytes)
        body = root.find(f'.//{_WNS}body') or root
        lines = []
        for para in body.iter(f'{_WNS}p'):
            texts = [t.text for t in para.iter(f'{_WNS}t') if t.text]
            line  = ''.join(texts).strip()
            if line:
                lines.append(line)
        return '\n'.join(lines)
    except Exception:
        return ''


def _extract_word_structured(fp: Path) -> dict:
    """
    Extract any Word-family file with structural analysis.

    .docx/.dotx/.docm/.dotm → python-docx (full structure)
    .doc/.dot               → try python-docx → OOXML ZIP → OLE2 stream
    """
    if not TDOC_DOCX_OK:
        return {'error': 'python-docx not installed'}

    fp  = Path(fp)
    suf = fp.suffix.lower()
    doc = None
    note = ''

    if suf in ('.docx', '.dotx', '.docm', '.dotm'):
        try:
            from docx import Document as _DocxDoc
            doc = _DocxDoc(str(fp))
        except Exception as e:
            return {'error': f'Cannot open {fp.name}: {e}'}

    elif suf in ('.doc', '.dot'):
        try:
            from docx import Document as _DocxDoc
            doc = _DocxDoc(str(fp))
        except Exception:
            pass

        if doc is None:
            text = _extract_ooxml_doc_text(fp)
            if text:
                refs = sorted(set(re.findall(r'[RSC]\d+-\d{6,7}', text)))
                return {'structured_text': text, 'metadata': {},
                        'section_count': 0, 'table_count': 0,
                        'char_count': len(text), 'references': refs,
                        'note': 'OOXML plain text (.doc extension)'}
            text, ole_note = _extract_ole2_doc_text(fp)
            if text:
                refs = sorted(set(re.findall(r'[RSC]\d+-\d{6,7}', text)))
                return {'structured_text': text, 'metadata': {},
                        'section_count': 0, 'table_count': 0,
                        'char_count': len(text), 'references': refs,
                        'note': ole_note}
            return {'error': f'Cannot extract text from {fp.name} ({ole_note})'}
    else:
        return {'error': f'Unsupported extension: {suf}'}

    if doc is None:
        return {'error': f'Failed to open {fp.name}'}

    # Full structural extraction via python-docx
    try:
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        meta    = _extract_3gpp_metadata(doc)
        sections = []
        cur      = {'heading': 'Document Start', 'level': 0, 'content': [], 'tables': []}
        table_counter = 0; last_para_text = None

        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = _para_text_with_revisions(para)
                if not text:
                    continue
                last_para_text = text
                style = para.style.name if para.style else ''
                if 'Heading' in style:
                    if cur['content'] or cur['tables']:
                        sections.append(cur)
                    lm  = re.search(r'Heading\s*(\d+)', style)
                    cur = {'heading': text,
                           'level':   int(lm.group(1)) if lm else 1,
                           'content': [], 'tables': []}
                else:
                    is_bold = any(r.bold for r in para.runs if r.text.strip())
                    if is_bold and len(text) < 100 and not cur['content']:
                        cur['content'].append(f'**{text}**')
                    else:
                        cur['content'].append(text)

            elif isinstance(element, CT_Tbl):
                table_counter += 1
                caption = _detect_table_caption(last_para_text, table_counter)
                cur['tables'].append(
                    _table_to_markdown(Table(element, doc), table_counter, caption))
                last_para_text = None

        if cur['content'] or cur['tables']:
            sections.append(cur)

        out_lines = []
        for sec in sections:
            if sec['level'] > 0:
                out_lines.append(f"\n{'#' * sec['level']} {sec['heading']}\n")
            if sec['content']:
                out_lines.append('\n'.join(sec['content']))
            if sec['tables']:
                out_lines.extend(sec['tables'])
            out_lines.append('')

        structured = '\n'.join(out_lines)
        full_text  = ' '.join(s['heading'] + ' ' + ' '.join(s['content']) for s in sections)
        refs       = sorted(set(re.findall(r'[RSC]\d+-\d{6,7}', full_text)))

        result = {
            'structured_text': structured,
            'metadata':        meta,
            'section_count':   len(sections),
            'table_count':     table_counter,
            'char_count':      len(structured),
            'references':      refs,
        }
        if note:
            result['note'] = note
        return result
    except Exception as e:
        return {'error': str(e)}


# ══════════════════════════════════════════════════════════════════════════════
# Public processing API
# ══════════════════════════════════════════════════════════════════════════════

def tdoc_process(uid: str, downloaded_paths, log_fn=None, out_dir=None, extract_dir=None) -> Path:
    """
    Extract all Word files (including inside ZIPs) và write all_tdocs_consolidated.md.

    Mỗi doc được enrich thêm ts_title (lookup từ ts_info.db qua TsInfoDb).

    Tier modes (based on number of docs with useful info):
      n <= 15          → full    (all fields incl. ts_title + headings)
      15 < n <= 30     → medium  (all fields incl. ts_title, no headings)
      n > 30           → compact (ts_number, ts_title, work_item, title, summary_of_change)

    Docs with no extractable information are silently excluded from the output file.

    extract_dir: thư mục chứa các file đã download/giải nén (do caller truyền vào tường minh).
                 Dùng để giải nén ZIP đúng chỗ, tránh đoán sai từ fp.parent.
    Returns output Path.
    """
    def log(m):
        if log_fn:
            log_fn(m)

    if out_dir is None:
        out_dir = DATA_DIR / 'outputs' / 'summary' / f'workitem_{uid}'
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    if extract_dir is not None:
        extract_dir = Path(extract_dir)

    # Step 1 — collect files, extract ZIPs
    all_files: list = []
    seen_zips: set  = set()
    for fp in downloaded_paths:
        fp = Path(fp)
        if not fp.exists():
            continue
        if fp.suffix.lower() == '.zip' and fp not in seen_zips:
            seen_zips.add(fp)
            stem     = fp.stem
            base_dir = extract_dir if extract_dir is not None else fp.parent
            xdir     = base_dir / stem
            if not xdir.exists():
                try:
                    xdir.mkdir(parents=True, exist_ok=True)
                    with zipfile.ZipFile(fp, 'r') as z:
                        z.extractall(xdir)
                    log(f"Extracted {fp.name} → {xdir}")
                except Exception as e:
                    log(f"[WARN] Zip error {fp.name}: {e}")
            all_files.extend(xdir.rglob('*') if xdir.exists() else [])
        else:
            all_files.append(fp)

    # Step 2 — unique Word files
    word_files = sorted(
        {f.resolve(): f for f in all_files
         if f.is_file() and f.suffix.lower() in _WORD_EXTS}.values()
    )
    exts_found = sorted({f.suffix.lower() for f in word_files})
    log(f"Xử lý {len(word_files)} Word file(s)"
        + (f" ({', '.join(exts_found)})" if exts_found else '') + "...")

    # Step 3 — load TS title lookup (once, before the loop)
    _ts_db = TsInfoDb(CACHE_DIR / "ts_info.db")

    # Step 3 — extract CR metadata + headings + ts_title
    processed = []
    for i, fp in enumerate(word_files, 1):
        log(f"  [{i}/{len(word_files)}] {fp.name}")

        # 3a — CR metadata (supports .docx + .doc)
        cr_meta = {}
        cr_error = None
        try:
            cr_meta = extract_cr_metadata(fp, output_txt=False)
        except Exception as e:
            cr_error = str(e)
            log(f"    ✗ CR metadata: {cr_error}")

        # 3b — TS title lookup from ts_info.db (O(1) dict lookup, no I/O)
        ts_number = cr_meta.get('ts_number') or ''
        ts_title  = _ts_db.get_title(ts_number) or ''

        # 3c — Headings (.docx only; skip .doc gracefully)
        headings: list[tuple[int, str]] = []
        heading_strategy = ''
        heading_error = None
        if fp.suffix.lower() in ('.docx', '.dotx', '.docm', '.dotm'):
            try:
                headings, heading_strategy = extract_headings_from_docx(fp)
                log(f"    → headings: {len(headings)} ({heading_strategy})")
            except Exception as e:
                heading_error = str(e)
                log(f"    ✗ headings: {heading_error}")
        else:
            heading_strategy = 'skipped (only .docx supported)'

        processed.append({
            'filename':                    fp.name,
            'tdoc_number':                 fp.stem,
            'ts_number':                   ts_number,
            'ts_title':                    ts_title,
            'work_item':                   cr_meta.get('work_item') or '',
            'title':                       cr_meta.get('title') or '',
            'reason_for_change':           cr_meta.get('reason_for_change') or '',
            'summary_of_change':           cr_meta.get('summary_of_change') or '',
            'consequences_if_not_approved': cr_meta.get('consequences_if_not_approved') or '',
            'other_comments':              cr_meta.get('other_comments') or '',
            'cr_error':                    cr_error,
            'headings':                    headings,
            'heading_strategy':            heading_strategy,
            'heading_error':               heading_error,
        })

    # Step 4 — filter out docs with no useful information
    # A doc is considered empty when ALL text fields are blank AND headings list is empty.
    _TEXT_FIELDS = [
        'ts_number', 'ts_title', 'work_item', 'title',
        'reason_for_change', 'summary_of_change',
        'consequences_if_not_approved', 'other_comments',
    ]

    def _has_info(doc: dict) -> bool:
        return any(doc.get(f) for f in _TEXT_FIELDS) or bool(doc.get('headings'))

    skipped_empty = [d for d in processed if not _has_info(d)]
    processed     = [d for d in processed if _has_info(d)]
    if skipped_empty:
        log(f"  ⚠ Bỏ qua {len(skipped_empty)} file không có thông tin: "
            + ", ".join(d['filename'] for d in skipped_empty))

    # Step 5 — write consolidated Markdown
    # Tier thresholds:
    #   n <= 15          → full    (all fields incl. ts_title + headings)
    #   15 < n <= 30     → medium  (all fields incl. ts_title, no headings)
    #   n > 30           → compact (ts_number, ts_title, work_item, title, summary_of_change)
    FULL_THRESHOLD   = 15
    MEDIUM_THRESHOLD = 30
    n = len(processed)
    if n <= FULL_THRESHOLD:
        mode = 'full'
    elif n <= MEDIUM_THRESHOLD:
        mode = 'medium'
    else:
        mode = 'compact'

    # ── Build unique TS summary (ordered by first appearance, deduped) ────────
    seen_ts: dict[str, str] = {}   # ts_number → ts_title (first occurrence wins)
    for doc in processed:
        ts = doc.get('ts_number') or ''
        if ts and ts not in seen_ts:
            seen_ts[ts] = doc.get('ts_title') or ''

    out_path = out_dir / 'all_tdocs_consolidated.md'
    with open(out_path, 'w', encoding='utf-8') as f:
        # ── File header ──────────────────────────────────────────────────────
        mode_suffix = {
            'full':    '',
            'medium':  ' *(Medium — All Fields, no Headings)*',
            'compact': ' *(Compact — Key Fields only)*',
        }[mode]
        f.write(f'# 3GPP TDoc CR Metadata — Consolidated{mode_suffix}\n\n')
        f.write('| | |\n|---|---|\n')
        f.write(f'| **Workitem** | `{uid}` |\n')
        f.write(f'| **Documents** | {n} |\n')
        if mode == 'medium':
            f.write(f'| **Mode** | Medium ({FULL_THRESHOLD} < docs ≤ {MEDIUM_THRESHOLD} — all fields, no headings) |\n')
        elif mode == 'compact':
            f.write(f'| **Mode** | Compact (> {MEDIUM_THRESHOLD} docs — TS Number, Work Item, Title, Summary of Change) |\n')
        f.write(f'| **Generated** | {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")} |\n\n')
        f.write('---\n\n')

        # ── TS Specifications referenced ─────────────────────────────────────
        if seen_ts:
            f.write('## TS Specifications Referenced\n\n')
            f.write('| TS Number | TS Title |\n')
            f.write('|---|---|\n')
            for ts_num, ts_ttl in seen_ts.items():
                title_cell = ts_ttl if ts_ttl else '*(not found)*'
                f.write(f'| `{ts_num}` | {title_cell} |\n')
            f.write('\n---\n\n')

        # ── Document index ────────────────────────────────────────────────────
        f.write('## Document Index\n\n')
        f.write('| # | TDoc | Title |\n')
        f.write('|---|---|---|\n')
        for idx, doc in enumerate(processed, 1):
            label = (doc['title'] or doc['filename'])[:80].replace('|', '\\|')
            f.write(f"| {idx} | `{doc['tdoc_number']}` | {label} |\n")
        f.write('\n---\n\n')

        # ── Per-document ──────────────────────────────────────────────────────
        f.write('## Documents\n\n')
        for idx, doc in enumerate(processed, 1):
            f.write(f"### {idx}. {doc['tdoc_number']}\n\n")

            if doc['cr_error']:
                f.write(f"> ⚠️ **CR metadata error:** {doc['cr_error']}\n\n")

            elif mode == 'compact':
                # ── Compact: ts_number, work_item, title, summary_of_change ──
                ts = doc.get('ts_number') or '—'
                wi = doc.get('work_item')  or '—'
                f.write(f'| TS Number | Work Item |\n|---|---|\n| `{ts}` | `{wi}` |\n\n')

                for key, label in [
                    ('title',             'Title'),
                    ('summary_of_change', 'Summary of Change'),
                ]:
                    value = doc.get(key) or '*(not found)*'
                    f.write(f'**{label}**\n\n{value}\n\n')

            elif mode == 'medium':
                # ── Medium: all fields, no headings ───────────────────────────
                ts = doc.get('ts_number') or '—'
                wi = doc.get('work_item')  or '—'
                f.write(f'| TS Number | Work Item |\n|---|---|\n| `{ts}` | `{wi}` |\n\n')

                for key, label in [
                    ('title',                        'Title'),
                    ('reason_for_change',            'Reason for Change'),
                    ('summary_of_change',            'Summary of Change'),
                    ('consequences_if_not_approved', 'Consequences if Not Approved'),
                    ('other_comments',               'Other Comments'),
                ]:
                    value = doc.get(key) or '*(not found)*'
                    f.write(f'**{label}**\n\n{value}\n\n')

            else:
                # ── Full: all fields + headings ───────────────────────────────
                ts = doc.get('ts_number') or '—'
                wi = doc.get('work_item')  or '—'
                f.write(f'| TS Number | Work Item |\n|---|---|\n| `{ts}` | `{wi}` |\n\n')

                for key, label in [
                    ('title',                        'Title'),
                    ('reason_for_change',            'Reason for Change'),
                    ('summary_of_change',            'Summary of Change'),
                    ('consequences_if_not_approved', 'Consequences if Not Approved'),
                    ('other_comments',               'Other Comments'),
                ]:
                    value = doc.get(key) or '*(not found)*'
                    f.write(f'**{label}**\n\n{value}\n\n')

                # ── Headings section ──────────────────────────────────────────
                h_label = (f"Headings ({doc['heading_strategy']})"
                           if doc['heading_strategy'] else 'Headings')
                f.write(f'#### {h_label}\n\n')
                if doc['heading_error']:
                    f.write(f'> ⚠️ Error: {doc["heading_error"]}\n\n')
                elif doc['headings']:
                    min_level = min(lvl for lvl, _ in doc['headings'])
                    for level, text in doc['headings']:
                        indent = '  ' * (level - min_level)
                        f.write(f'{indent}- {text}\n')
                    f.write('\n')
                else:
                    f.write('*(none found)*\n\n')

            f.write('---\n\n')

    log(f"Processed {n} file(s) [{mode}] → {out_path}")
    return out_path
